"""
doc_utils.py  — Shared formatting utilities for Vishal Jewelers client document generation.
Aptos font, 12pt body, justified paragraphs, professional hierarchy.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime
import os


def rgb_to_hex(color) -> str:
    """Convert an RGBColor (r, g, b) tuple to a 6-character uppercase hex string."""
    r, g, b = color[0], color[1], color[2]
    return f"{r:02X}{g:02X}{b:02X}"

FONT_BODY     = "Aptos"
FONT_HEADING  = "Aptos"
SIZE_BODY     = 12
SIZE_H1       = 20
SIZE_H2       = 15
SIZE_H3       = 13
SIZE_SMALL    = 10

COLOR_DARK    = RGBColor(0x1A, 0x1A, 0x2E)   # Dark navy
COLOR_GOLD    = RGBColor(0xD4, 0xAF, 0x37)   # Metallic Gold (brand)
COLOR_MID     = RGBColor(0x44, 0x44, 0x44)   # Body text grey
COLOR_WHITE   = RGBColor(0xFF, 0xFF, 0xFF)
COLOR_LIGHT   = RGBColor(0xF5, 0xF5, 0xF5)


def new_document() -> Document:
    """Create a new document with standard margins."""
    doc = Document()
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.5)
    return doc


def set_run_font(run, name=FONT_BODY, size=SIZE_BODY, bold=False, italic=False, color=None):
    run.font.name  = name
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    # Ensure Aptos is applied at Complex-Script level too
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"),       name)
    rFonts.set(qn("w:hAnsi"),       name)
    rFonts.set(qn("w:cs"),          name)
    rFonts.set(qn("w:eastAsia"),    name)
    existing = rPr.find(qn("w:rFonts"))
    if existing is not None:
        rPr.remove(existing)
    rPr.insert(0, rFonts)


def add_cover(doc: Document, title: str, subtitle: str, version: str = "v1.0",
              date: str = None, prepared_for: str = "Vishal Jewelers",
              prepared_by: str = "Engineering Team"):
    """Insert a professional cover page."""
    if date is None:
        date = datetime.date.today().strftime("%B %d, %Y")

    # Top spacer
    for _ in range(6):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)

    # Document title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(title.upper())
    set_run_font(run, size=SIZE_H1, bold=True, color=COLOR_DARK)

    # Gold rule
    add_horizontal_rule(doc, color=COLOR_GOLD)

    # Subtitle
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(subtitle)
    set_run_font(run, size=14, italic=True, color=COLOR_MID)

    for _ in range(3):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)

    # Metadata table
    table = doc.add_table(rows=5, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    meta = [
        ("Project",      "Vishal Jewelers — B2B/B2C Jewelry App"),
        ("Prepared For", prepared_for),
        ("Prepared By",  prepared_by),
        ("Version",      version),
        ("Date",         date),
    ]
    for i, (label, value) in enumerate(meta):
        row = table.rows[i]
        lc  = row.cells[0]
        vc  = row.cells[1]
        lc.width = Inches(1.8)
        vc.width = Inches(3.5)
        shade_cell(lc, "1A1A2E")
        lp = lc.paragraphs[0]
        lr = lp.add_run(label)
        set_run_font(lr, bold=True, color=COLOR_WHITE, size=11)
        vp = vc.paragraphs[0]
        vr = vp.add_run(value)
        set_run_font(vr, size=11)

    doc.add_page_break()


def shade_cell(cell, hex_color: str):
    """Fill a table cell with a solid background color."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def add_horizontal_rule(doc: Document, color: RGBColor = COLOR_GOLD, thickness_pt: int = 2):
    """Add a coloured horizontal rule paragraph."""
    p  = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after  = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    str(thickness_pt * 8))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), rgb_to_hex(color))
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_h1(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    set_run_font(run, size=SIZE_H1, bold=True, color=COLOR_DARK)
    add_horizontal_rule(doc, color=COLOR_GOLD, thickness_pt=1)
    return p


def add_h2(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    set_run_font(run, size=SIZE_H2, bold=True, color=COLOR_DARK)
    return p


def add_h3(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    set_run_font(run, size=SIZE_H3, bold=True, color=COLOR_MID)
    return p


def add_body(doc: Document, text: str, justify: bool = True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if justify else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after  = Pt(8)
    p.paragraph_format.first_line_indent = Pt(0)
    run = p.add_run(text)
    set_run_font(run, size=SIZE_BODY, color=COLOR_MID)
    return p


def add_bullet(doc: Document, text: str, level: int = 0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.3 + level * 0.3)
    run = p.add_run(text)
    set_run_font(run, size=SIZE_BODY, color=COLOR_MID)
    return p


def add_numbered(doc: Document, text: str):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_run_font(run, size=SIZE_BODY, color=COLOR_MID)
    return p


def add_note(doc: Document, label: str, text: str):
    """Add a highlighted note/callout box."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r1 = p.add_run(f"{label}: ")
    set_run_font(r1, bold=True, size=SIZE_BODY, color=COLOR_GOLD)
    r2 = p.add_run(text)
    set_run_font(r2, size=SIZE_BODY, color=COLOR_MID)


def add_table(doc: Document, headers: list, rows: list):
    """Add a styled table with gold header row."""
    col_count = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=col_count)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hrow = table.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        shade_cell(cell, "1A1A2E")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_run_font(run, bold=True, size=11, color=COLOR_WHITE)

    # Data rows
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        bg  = "F5F5F5" if ri % 2 == 0 else "FFFFFF"
        for ci, cell_text in enumerate(row_data):
            cell = row.cells[ci]
            shade_cell(cell, bg)
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            set_run_font(run, size=10, color=COLOR_MID)

    doc.add_paragraph()   # spacing after table


def add_signature_block(doc: Document, parties: list):
    """parties = [('Client Representative', 'Vishal Jewelers'), ('Project Manager', 'Our Firm')]"""
    doc.add_paragraph()
    add_h2(doc, "Authorization & Sign-Off")
    add_body(doc, "The undersigned parties confirm that they have reviewed the contents of this document and agree to its terms and findings.")
    doc.add_paragraph()

    for role, org in parties:
        table = doc.add_table(rows=3, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        labels = [("Name:", ""), ("Signature:", ""), ("Date:", "")]
        for ri, (lbl, _) in enumerate(labels):
            c0 = table.rows[ri].cells[0]
            c1 = table.rows[ri].cells[1]
            r = c0.paragraphs[0].add_run(lbl)
            set_run_font(r, bold=True, size=10)
            r2 = c1.paragraphs[0].add_run("_" * 40)
            set_run_font(r2, size=10, color=COLOR_MID)

        p = doc.add_paragraph()
        r = p.add_run(f"{role}  |  {org}")
        set_run_font(r, italic=True, size=SIZE_SMALL, color=COLOR_MID)
        doc.add_paragraph()


def add_toc_placeholder(doc: Document, entries: list):
    """Add a manual table of contents."""
    add_h1(doc, "Table of Contents")
    for i, (section, page_ref) in enumerate(entries, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(f"{i}. {section}")
        set_run_font(run, size=SIZE_BODY, color=COLOR_MID)
    doc.add_page_break()


def add_screenshot(doc: Document, image_path: str, caption: str = "", width_inches: float = 2.5):
    """Insert a centred screenshot with an optional italic caption."""
    if not os.path.exists(image_path):
        return  # silently skip missing images
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(0)
    run = p.add_run()
    run.add_picture(image_path, width=Inches(width_inches))
    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_before = Pt(2)
        cp.paragraph_format.space_after  = Pt(10)
        cr = cp.add_run(caption)
        set_run_font(cr, size=9, italic=True, color=COLOR_MID)


def save(doc: Document, path: str):
    doc.save(path)
    print(f"  ✓ Saved: {path}")
